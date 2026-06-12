from __future__ import annotations

from pathlib import Path

from mapemgen.ingestion.fact_records import make_fact
from mapemgen.ingestion.movement_tables import extract_utc_form_movement_facts
from mapemgen.ingestion.text_facts import extract_keyword_facts, extract_metadata_facts


def extract_docx_facts(path: str | Path) -> list[dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX extraction requires the 'python-docx' package.") from exc

    document = Document(path)
    name = Path(path).name.lower()
    source_role = "utc_form" if "utc" in name or "utcform" in name else None
    facts: list[dict] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = " ".join(paragraph.text.split())
        if text:
            facts.extend(extract_keyword_facts([text], f"paragraph {index}", exact_location=True, source_role=source_role))
            facts.extend(extract_metadata_facts([text], f"paragraph {index}", exact_location=True))
    table_matrices: list[list[list[str]]] = []
    for table_index, table in enumerate(document.tables, start=1):
        table_matrix: list[list[str]] = []
        for row_index, row in enumerate(table.rows, start=1):
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
            table_matrix.append(cells)
            if not cells:
                continue
            value = " | ".join(cells)
            location = f"table {table_index} row {row_index}"
            facts.append(_fact("docx_table_row", value, location, 0.8))
            facts.extend(extract_keyword_facts([value], location, exact_location=True, source_role=source_role))
            facts.extend(extract_metadata_facts([value], location, exact_location=True))
        table_matrices.append(table_matrix)
    if source_role == "utc_form":
        facts.extend(extract_utc_form_movement_facts(table_matrices))
    return facts


def _fact(fact_name: str, value: object, location: str, confidence: float) -> dict:
    return make_fact(fact_name, value, location, confidence)
